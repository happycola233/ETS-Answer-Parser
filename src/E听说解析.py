# pip install python-docx

import os
import json
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE

# 获取用户的桌面路径
desktop_path = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')

# 获取%appdata%\ETS路径
appdata_path = os.getenv('APPDATA')
ets_path = os.path.join(appdata_path, 'ETS')

# 获取按名称排序的试卷子文件夹
def get_sorted_content_folders(folder_path):
    # 获取子文件夹，并按照名称的数字部分进行排序
    content_folders = [f for f in os.listdir(folder_path) if f.startswith('content_')]
    content_folders.sort(key=lambda x: int(x.split('_')[1]))  # 按数字部分排序
    return content_folders

# 去除HTML标签的函数，处理换行
def clean_html(raw_html):
    # 将<p>和<br>替换为特殊标记，以避免连续的换行符
    raw_html = raw_html.replace('<p>', '[NEWLINE]').replace('</p>', '[NEWLINE]')
    raw_html = raw_html.replace('<br>', '[NEWLINE]').replace('<br/>', '[NEWLINE]').replace('</br>', '[NEWLINE]')
    # 去除其他HTML标签
    clean_text = re.sub('<.*?>', '', raw_html)
    # 将特殊标记替换为单个换行符
    clean_text = clean_text.replace('[NEWLINE]', '\n')
    # 去除多余的空行
    clean_text = re.sub(r'\n+', '\n', clean_text).strip()
    return clean_text

# 解析某个文件夹中的Section
def parse_section_from_folder(folder_path, content_folder_name, start_number):
    content_folder = os.path.join(folder_path, content_folder_name)
    json_file_path = os.path.join(content_folder, 'content2.json')

    # 检查文件是否存在
    if not os.path.exists(json_file_path):
        print(f"文件 {json_file_path} 不存在，跳过此文件夹。")
        return "", start_number

    # 读取JSON文件
    with open(json_file_path, 'r', encoding='utf-8') as file:
        data = json.load(file)

    questions = data['info']['xtlist']
    parsed_questions = []

    # 解析每一道题目
    for question in questions:
        # 去除题目中的HTML标签
        question_text = clean_html(question['xt_value'])

        # 检查题目中是否包含题号，避免重复
        if not question_text.strip().startswith(str(start_number) + '.'):
            question_text = f"{start_number}. {question_text}"

        options = question['xxlist']
        answers = "\n".join([f"{opt['xx_mc']}. {clean_html(opt['xx_nr'])}" for opt in options])
        correct_answer = question['answer']
        # 在每个小题之间添加一个空行
        parsed_questions.append(f"{question_text}\n{answers}\n答案：{correct_answer}\n\n")
        start_number += 1  # 题目序号递增

    return "\n".join(parsed_questions), start_number

# 解析Section A
def parse_section_a(folder_path):
    content_folders = get_sorted_content_folders(folder_path)
    start_number = 1
    section_a_content = ""
    section_a_part1, start_number = parse_section_from_folder(folder_path, content_folders[0], start_number)
    section_a_content += section_a_part1
    section_a_part2, start_number = parse_section_from_folder(folder_path, content_folders[1], start_number)
    section_a_content += section_a_part2
    return section_a_content

# 解析Section B
def parse_section_b_with_reading(folder_path):
    content_folders = get_sorted_content_folders(folder_path)
    start_number = 11
    section_b_content = ""
    for i in range(2, 5):
        content_folder = os.path.join(folder_path, content_folders[i])
        json_file_path = os.path.join(content_folder, 'content2.json')
        if not os.path.exists(json_file_path):
            print(f"文件 {json_file_path} 不存在，跳过此文件夹。")
            continue
        with open(json_file_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
        reading_material = clean_html(data['info']['st_nr'])
        section_b_content += f"{reading_material}\n\n"
        questions = data['info']['xtlist']
        for question in questions:
            question_text = clean_html(question['xt_value'])
            if not question_text.strip().startswith(str(start_number) + '.'):
                question_text = f"{start_number}. {question_text}"
            options = question['xxlist']
            answers = "\n".join([f"{opt['xx_mc']}. {clean_html(opt['xx_nr'])}" for opt in options])
            correct_answer = question['answer']
            # 在每个小题之间添加一个空行
            section_b_content += f"{question_text}\n{answers}\n答案：{correct_answer}\n\n"
            start_number += 1
    return section_b_content

# 解析朗读句子
def parse_read_sentences(folder_path):
    content_folders = get_sorted_content_folders(folder_path)
    read_sentences_content = ""
    for i in range(5, 7):
        content_folder = os.path.join(folder_path, content_folders[i])
        json_file_path = os.path.join(content_folder, 'content2.json')
        with open(json_file_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
        read_sentence = clean_html(data['info']['value'])
        read_sentences_content += f"{read_sentence}\n\n"
    return read_sentences_content

# 解析朗读段落
def parse_read_paragraph(folder_path):
    content_folders = get_sorted_content_folders(folder_path)
    content_folder = os.path.join(folder_path, content_folders[7])
    json_file_path = os.path.join(content_folder, 'content2.json')
    with open(json_file_path, 'r', encoding='utf-8') as file:
        data = json.load(file)
    read_paragraph = clean_html(data['info']['value'])
    return read_paragraph

# 解析情景提问
def parse_scenario_questions(folder_path):
    content_folders = get_sorted_content_folders(folder_path)
    content_folder = os.path.join(folder_path, content_folders[8])
    json_file_path = os.path.join(content_folder, 'content2.json')
    with open(json_file_path, 'r', encoding='utf-8') as file:
        data = json.load(file)
    scenario_questions_content = ""
    for question_info in data['info']['question']:
        ask_text = clean_html(question_info['ask'])
        scenario_questions_content += f"{ask_text}\n答案:\n"
        for std in question_info['std']:
            scenario_questions_content += f"● {std['value']}\n"
        scenario_questions_content += "\n"
    return scenario_questions_content

# 解析图片描述
def parse_picture_scenario(folder_path):
    content_folders = get_sorted_content_folders(folder_path)
    content_folder = os.path.join(folder_path, content_folders[9])  # 第10个文件夹
    json_file_path = os.path.join(content_folder, 'content2.json')

    with open(json_file_path, 'r', encoding='utf-8') as file:
        data = json.load(file)

    picture_scenario_content = "答案:\n"

    for std in data['info']['std']:
        picture_scenario_content += f"● {std['value']}\n"

    # 解析并格式化keypoint
    keypoint = clean_html(data['info']['keypoint'])
    picture_scenario_content += f"\nkeypoint:\n{keypoint}\n"

    # 直接在material文件夹中寻找content.jpg
    material_folder = os.path.join(content_folder, 'material')
    image_path = os.path.join(material_folder, 'content.jpg')

    if not os.path.exists(image_path):
        print(f"图片 {image_path} 不存在。")
        image_path = None

    return picture_scenario_content, image_path

# 解析快速应答
def parse_quick_response(folder_path):
    content_folders = get_sorted_content_folders(folder_path)
    content_folder = os.path.join(folder_path, content_folders[10])  # 第11个文件夹
    json_file_path = os.path.join(content_folder, 'content2.json')

    quick_response_content = ""
    with open(json_file_path, 'r', encoding='utf-8') as file:
        data = json.load(file)

    questions = data['info']['question']
    for question_info in questions:
        # 题目
        ask = clean_html(question_info['ask'])
        quick_response_content += f"{ask}\n"

        # 答案
        quick_response_content += "答案:\n"
        for std in question_info['std']:
            quick_response_content += f"● {std['value']}\n"

        # 关键词
        keywords = question_info.get('keywords', '')
        if keywords:
            quick_response_content += f"\n关键词:\n{keywords}\n"

        quick_response_content += "\n"
    return quick_response_content

# 解析简述和回答
def parse_summary_and_answer(folder_path):
    content_folders = get_sorted_content_folders(folder_path)
    content_folder = os.path.join(folder_path, content_folders[11])  # 第12个文件夹
    json_file_path = os.path.join(content_folder, 'content2.json')

    summary_and_answer_content = ""
    with open(json_file_path, 'r', encoding='utf-8') as file:
        data = json.load(file)

    # 提取原文
    original_text = clean_html(data['info']['value'])
    summary_and_answer_content += f"原文：\n{original_text}\n\n"

    # 提取每个问题和答案
    questions = data['info']['question']
    for question_info in questions:
        ask = clean_html(question_info['ask'])
        summary_and_answer_content += f"{ask}\n答案:\n"
        for std in question_info['std']:
            summary_and_answer_content += f"● {std['value']}\n"
        summary_and_answer_content += "\n"
    return summary_and_answer_content

# 获取唯一的文件名，避免覆盖
def get_unique_filename(filename):
    base, ext = os.path.splitext(filename)
    counter = 1
    new_filename = filename
    while os.path.exists(new_filename):
        new_filename = f"{base}_{counter}{ext}"
        counter += 1
    return new_filename

# 保存解析结果到Word文档
def save_to_word(content, filename):
    document = Document()

    # 设置默认字体为等线，字号为12磅
    style = document.styles['Normal']
    font = style.font
    font.name = '等线'
    font.size = Pt(12)
    font.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
    font.element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    font.element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    font.element.rPr.rFonts.set(qn('w:cs'), 'Times New Roman')

    # 设置段落格式：单倍行距，段前段后间距为0
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # 添加各部分到Word文档
    def add_content_with_style(title, text, section_name=None):
        heading = document.add_heading(title, level=1)
        heading.bold = True
        # 设置大标题字体
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
            # 设置大标题颜色为绿色
            run.font.color.rgb = RGBColor(0, 176, 80)  # 绿色 #00B050

        # 将文本按照双换行符拆分为段落
        paragraphs = re.split(r'\n{2,}', text)
        for para in paragraphs:
            lines = para.strip().split('\n')
            for idx, line in enumerate(lines):
                if not line.strip():
                    continue  # 跳过空行
                if section_name in ['section_a', 'section_b'] and line.startswith('答案：'):
                    # 对于Section A和B，将'答案：'加粗，答案设为蓝色
                    p = document.add_paragraph()
                    # '答案：'部分
                    run = p.add_run('答案：')
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
                    # 答案部分
                    answer_text = line[len('答案：'):]
                    run = p.add_run(answer_text)
                    run.font.color.rgb = RGBColor(0, 112, 192)
                    run.font.name = 'Times New Roman'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
                    # 设置段落格式
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                elif line.startswith('答案：') or line == '答案:' or line.startswith('keypoint:') or line == '原文：':
                    p = document.add_paragraph()
                    run = p.add_run(line)
                    run.bold = True
                    # 设置字体
                    run.font.name = 'Times New Roman'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
                    # 设置段落格式
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                elif line.startswith('● '):
                    p = document.add_paragraph()
                    run = p.add_run(line)
                    # 设置字体颜色为蓝色
                    run.font.color.rgb = RGBColor(0, 112, 192)  # 蓝色
                    # 设置字体
                    run.font.name = 'Times New Roman'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
                    # 设置段落格式
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                else:
                    p = document.add_paragraph()
                    run = p.add_run(line)
                    # 设置字体
                    run.font.name = 'Times New Roman'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
                    # 设置段落格式
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            # 在每个小题之间添加一个空行
            document.add_paragraph()

    add_content_with_style('Section A', content['section_a'], 'section_a')
    add_content_with_style('Section B', content['section_b'], 'section_b')
    add_content_with_style('朗读句子', content['read_sentences'])
    add_content_with_style('朗读段落', content['read_paragraph'])
    add_content_with_style('情景提问', content['scenario_questions'])

    heading = document.add_heading('图片描述', level=1)
    heading.bold = True
    # 设置大标题字体
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
        # 设置大标题颜色为绿色
        run.font.color.rgb = RGBColor(0, 176, 80)

    # 插入图片
    if content['picture_scenario_image']:
        document.add_picture(content['picture_scenario_image'], width=Inches(4))
    else:
        document.add_paragraph('（图片缺失）')

    # 添加答案和keypoint
    paragraphs = re.split(r'\n{2,}', content['picture_scenario'])
    for para in paragraphs:
        lines = para.strip().split('\n')
        for idx, line in enumerate(lines):
            if not line.strip():
                continue  # 跳过空行
            if line == '答案:' or line.startswith('keypoint:') or line == '原文：':
                p = document.add_paragraph()
                run = p.add_run(line)
                run.bold = True
            elif line.startswith('● '):
                p = document.add_paragraph()
                run = p.add_run(line)
                # 设置字体颜色为蓝色
                run.font.color.rgb = RGBColor(0, 112, 192)
            else:
                p = document.add_paragraph()
                run = p.add_run(line)
            # 设置字体
            run.font.name = 'Times New Roman'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
            # 设置段落格式
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    add_content_with_style('快速应答', content['quick_response'])
    add_content_with_style('简述和回答', content['summary_and_answer'])

    # 保存文档，避免覆盖
    filename = get_unique_filename(filename)
    document.save(filename)
    return filename  # 返回最终保存的文件名

# 主程序
if __name__ == "__main__":
    # 获取%appdata%\ETS路径下按创建时间排序的试卷文件夹
    folders = [(f, os.path.getctime(os.path.join(ets_path, f))) for f in os.listdir(ets_path) if os.path.isdir(os.path.join(ets_path, f))]
    folders.sort(key=lambda x: x[1], reverse=True)  # 按照创建时间降序排序（新到旧）
    sorted_folders = [f[0] for f in folders]

    # 用户选择试卷文件夹
    print("请选择试卷文件夹（已按照下载时间从新到旧排列）：")
    for i, folder in enumerate(sorted_folders):
        print(f"{i + 1}. {folder}")
    choice = int(input("请输入对应的数字: ")) - 1
    chosen_folder = sorted_folders[choice]
    chosen_folder_path = os.path.join(ets_path, chosen_folder)

    # 解析各部分内容
    section_a_content = parse_section_a(chosen_folder_path)
    section_b_content = parse_section_b_with_reading(chosen_folder_path)
    read_sentences_content = parse_read_sentences(chosen_folder_path)
    read_paragraph_content = parse_read_paragraph(chosen_folder_path)
    scenario_questions_content = parse_scenario_questions(chosen_folder_path)
    picture_scenario_content, picture_image_path = parse_picture_scenario(chosen_folder_path)
    quick_response_content = parse_quick_response(chosen_folder_path)
    summary_and_answer_content = parse_summary_and_answer(chosen_folder_path)

    # 保存到Word文档
    content = {
        'section_a': section_a_content,
        'section_b': section_b_content,
        'read_sentences': read_sentences_content,
        'read_paragraph': read_paragraph_content,
        'scenario_questions': scenario_questions_content,
        'picture_scenario': picture_scenario_content,
        'picture_scenario_image': picture_image_path,
        'quick_response': quick_response_content,
        'summary_and_answer': summary_and_answer_content
    }
    word_file = os.path.join(desktop_path, 'E听说_解析.docx')
    final_filename = save_to_word(content, word_file)

    print(f"解析完成，文件已保存到: {final_filename}")
